#!/usr/bin/env python3
"""
sync_granola.py — mirror Granola notes + transcripts to local .docx files

Compatible with Granola v7.195+ which moved local storage to encrypted SQLite.
Data is fetched from the Granola cloud API using the WorkOS bearer token stored
in supabase.json.enc (or the legacy plaintext supabase.json).

Usage:
    python sync_granola.py                # incremental sync
    python sync_granola.py --dry-run      # no .docx writes
    python sync_granola.py --full         # ignore manifest, rebuild everything
    python sync_granola.py --folder PMF   # restrict to one folder (substring match)
    python sync_granola.py --skip-transcripts  # skip transcript fetches (faster)
    python sync_granola.py --healthcheck  # verify deps/auth/API/filesystem
    python sync_granola.py --reconcile    # compare API, manifest, docx
"""
from __future__ import annotations

import argparse
import base64
import gzip
import json
import logging
import os
import re
import shutil
import socket
import subprocess
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Paths and constants
# ---------------------------------------------------------------------------

PROJECT_DIR = Path(__file__).resolve().parent
STATE_DIR   = Path(os.environ.get(
    "GRANOLA_SYNC_STATE_DIR",
    Path.home() / "Library" / "Application Support" / "granola-sync",
))
BUILD_DIR   = Path(os.environ.get(
    "GRANOLA_SYNC_OUTPUT",
    Path.home() / "Documents" / "Granola Notes",
))
LOGS_DIR      = STATE_DIR / "logs"
MANIFEST_PATH = STATE_DIR / "manifest.json"

GRANOLA_SUPPORT = Path.home() / "Library" / "Application Support" / "Granola"
TOKEN_FILE      = GRANOLA_SUPPORT / "supabase.json"        # legacy plaintext (may be stale)
SUPABASE_ENC_FILE = GRANOLA_SUPPORT / "supabase.json.enc"  # current encrypted token
DEK_FILE        = GRANOLA_SUPPORT / "storage.dek"          # encrypted DEK (Chromium safeStorage)

GRANOLA_API = "https://api.granola.ai"
GRANOLA_CLIENT_VERSION = "7.195.0"
UNFILED_FOLDER = "_Unfiled"

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

def setup_logging() -> logging.Logger:
    LOGS_DIR.mkdir(exist_ok=True)
    log_path = LOGS_DIR / f"sync-{datetime.now().strftime('%Y-%m-%d')}.log"
    logger = logging.getLogger("sync_granola")
    logger.setLevel(logging.INFO)
    fmt = logging.Formatter("%(asctime)s %(levelname)s %(message)s")
    fh = logging.FileHandler(log_path)
    fh.setFormatter(fmt)
    sh = logging.StreamHandler(sys.stdout)
    sh.setFormatter(fmt)
    logger.handlers = [fh, sh]
    return logger

# ---------------------------------------------------------------------------
# Granola client — cloud API
# ---------------------------------------------------------------------------

class GranolaClient:
    """Authenticates with Granola's cloud API and fetches notes data.

    Token source: supabase.json.enc (AES-256-GCM, decrypted via Keychain + storage.dek).
    Falls back to the legacy plaintext supabase.json if decryption fails.
    """

    def __init__(self, logger: logging.Logger):
        self.log = logger

    # ---- DEK / encryption helpers ----

    def _get_dek(self) -> bytes | None:
        """Derive the 32-byte storage DEK from the macOS Keychain + storage.dek.

        Granola uses Chromium's safeStorage (v10 CBC format):
          1. Keychain password → PBKDF2-HMAC-SHA1(pw, salt=b"saltysalt", iter=1003, len=16)
          2. AES-128-CBC(key, iv=b" "*16) decrypt storage.dek[3:]
          3. PKCS7 unpad → base64 decode → 32-byte DEK
        """
        if not DEK_FILE.exists():
            return None
        try:
            from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
            from cryptography.hazmat.primitives import hashes
            from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
            from cryptography.hazmat.backends import default_backend

            res = subprocess.run(
                ["security", "find-generic-password",
                 "-s", "Granola Safe Storage", "-a", "Granola Key", "-w"],
                capture_output=True, text=True, timeout=5,
            )
            if res.returncode != 0 or not res.stdout.strip():
                return None

            pw = res.stdout.strip().encode()
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA1(), length=16,
                salt=b"saltysalt", iterations=1003, backend=default_backend(),
            )
            aes_key = kdf.derive(pw)

            enc = DEK_FILE.read_bytes()  # b"v10" + 48-byte AES-128-CBC ciphertext
            decryptor = Cipher(
                algorithms.AES(aes_key), modes.CBC(b" " * 16), backend=default_backend()
            ).decryptor()
            raw = decryptor.update(enc[3:]) + decryptor.finalize()
            pad = raw[-1]
            dek = base64.b64decode(raw[:-pad])
            return dek if len(dek) == 32 else None
        except Exception as e:
            self.log.debug(f"  _get_dek failed: {e}")
            return None

    def _decrypt_enc_file(self, path: Path, dek: bytes) -> dict:
        """Decrypt a Granola .enc file (AES-256-GCM, 12-byte nonce prefix)."""
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
        data = path.read_bytes()
        plaintext = AESGCM(dek).decrypt(data[:12], data[12:], None)
        return json.loads(plaintext)

    # ---- authentication ----

    def _token(self) -> str:
        """Return a valid WorkOS bearer token.

        Tries supabase.json.enc (current) first; falls back to plaintext supabase.json.
        """
        if SUPABASE_ENC_FILE.exists():
            dek = self._get_dek()
            if dek:
                try:
                    supabase = self._decrypt_enc_file(SUPABASE_ENC_FILE, dek)
                    workos_raw = supabase.get("workos_tokens", "{}")
                    workos = json.loads(workos_raw) if isinstance(workos_raw, str) else workos_raw
                    return workos["access_token"]
                except Exception as e:
                    self.log.debug(f"  supabase.json.enc decrypt failed: {e}")

        # Fallback: legacy plaintext (may be stale in Granola v7.195+)
        if not TOKEN_FILE.exists():
            raise SystemExit(
                "Granola token not found. Open the Granola desktop app and try again."
            )
        with TOKEN_FILE.open() as f:
            outer = json.load(f)
        try:
            return json.loads(outer["workos_tokens"])["access_token"]
        except (KeyError, json.JSONDecodeError) as e:
            raise SystemExit(f"Could not parse Granola token from {TOKEN_FILE}: {e}")

    # ---- HTTP ----

    def _post(self, path: str, body: dict[str, Any], timeout: int = 60, retries: int = 2) -> Any:
        headers = {
            "Authorization": f"Bearer {self._token()}",
            "Content-Type": "application/json",
            "Accept": "*/*",
            "Accept-Encoding": "gzip",
            "User-Agent": f"Granola/{GRANOLA_CLIENT_VERSION}",
        }
        body = {**body, "clientVersion": GRANOLA_CLIENT_VERSION}
        req = urllib.request.Request(
            f"{GRANOLA_API}{path}",
            data=json.dumps(body).encode(),
            headers=headers,
            method="POST",
        )
        for attempt in range(retries + 1):
            try:
                with urllib.request.urlopen(req, timeout=timeout) as r:
                    raw = r.read()
                    if r.headers.get("Content-Encoding") == "gzip":
                        raw = gzip.decompress(raw)
                    return json.loads(raw.decode())
            except urllib.error.HTTPError as e:
                if e.code == 401:
                    raise SystemExit(
                        "Granola API returned 401. Token expired — open the Granola desktop "
                        "app (which refreshes the token) and re-run."
                    )
                if attempt == retries:
                    # Raise a plain exception so per-doc callers can catch and skip
                    raise RuntimeError(f"Granola API {path} HTTP {e.code}: {e.read()[:200].decode(errors='replace')}")
                self.log.warning(f"  {path} HTTP {e.code}, retrying ({attempt+1}/{retries})")
            except (urllib.error.URLError, socket.timeout) as e:
                if attempt == retries:
                    # Raise a plain exception so per-doc callers can catch and skip
                    raise RuntimeError(f"Granola API {path} network error: {e}")
                self.log.warning(f"  {path} network error: {e}, retrying ({attempt+1}/{retries})")
            time.sleep(1.5 * (attempt + 1))

    # ---- folders ----

    def get_all_document_lists_metadata(self) -> dict[str, dict]:
        """Return {folder_id: {title, deleted_at, ...}} from the cloud API."""
        try:
            result = self._post("/v1/get-document-lists-metadata", {})
        except RuntimeError as e:
            raise SystemExit(f"Could not fetch folder list: {e}")
        return result.get("lists", {}) if isinstance(result, dict) else {}

    def get_document_list_with_docs(self, list_id: str) -> dict:
        """Return folder metadata + embedded document list from the cloud API.

        The 'documents' key in the response is a list of full document objects
        (title, created_at, updated_at, people, google_calendar_event, etc.).
        """
        try:
            return self._post("/v1/get-document-list", {"list_id": list_id}, timeout=30)
        except RuntimeError as e:
            raise SystemExit(f"Could not fetch folder contents: {e}")

    def get_recent_documents(self) -> list[dict]:
        """Return recently touched documents regardless of folder membership.

        Catches documents that aren't filed into any Granola folder (shown as
        owner "Me" with no folder in the app). Not a full history — the API
        returns a rolling window of recently created/updated docs, not true
        pagination — but it's the only way to surface unfiled notes at all.
        """
        try:
            result = self._post("/v1/get-documents", {}, timeout=30)
            return result if isinstance(result, list) else []
        except Exception as e:
            self.log.warning(f"  could not fetch recent/unfiled documents: {e}")
            return []

    # ---- panels (AI-generated summaries) ----

    def get_document_panels(self, document_id: str) -> list[dict[str, Any]]:
        """Return AI summary panels for a document."""
        d = self._post("/v1/get-document-panels", {"document_id": document_id}, timeout=30)
        if isinstance(d, list):
            return d
        raise ValueError(f"Unexpected panels response for {document_id}: {type(d).__name__}")

    # ---- transcript ----

    def get_transcript(self, document_id: str) -> list[dict[str, Any]]:
        """Return transcript segments from the cloud API."""
        d = self._post("/v1/get-document-transcript", {"document_id": document_id}, timeout=30)
        if isinstance(d, list):
            return d
        raise ValueError(f"Unexpected transcript response for {document_id}: {type(d).__name__}")

# ---------------------------------------------------------------------------
# Folder + document indexing from API
# ---------------------------------------------------------------------------

@dataclass
class FolderIndex:
    """folder_id → title; doc_id → folder_title; doc_id → doc object."""
    id_to_title:  dict[str, str]        = field(default_factory=dict)
    doc_to_folder: dict[str, str]       = field(default_factory=dict)
    documents:    dict[str, dict]       = field(default_factory=dict)


def build_folder_index(client: GranolaClient, log: logging.Logger) -> FolderIndex:
    """Fetch all folders and their documents from the Granola cloud API.

    Makes one call to get-document-lists-metadata for folder names,
    then one call to get-document-list per folder for the document objects.
    """
    idx = FolderIndex()

    # 1. All folder metadata
    metadata = client.get_all_document_lists_metadata()
    for fid, meta in metadata.items():
        if not isinstance(meta, dict) or meta.get("deleted_at"):
            continue
        idx.id_to_title[fid] = meta.get("title") or "Untitled Folder"

    log.info(f"  fetched metadata for {len(idx.id_to_title)} folders")

    # 2. Per-folder document objects
    for fid, title in idx.id_to_title.items():
        try:
            folder_data = client.get_document_list_with_docs(fid)
            docs = folder_data.get("documents", [])
            doc_list = docs if isinstance(docs, list) else []
            for doc in doc_list:
                if not isinstance(doc, dict):
                    continue
                doc_id = doc.get("id")
                if not doc_id:
                    continue
                if doc_id not in idx.doc_to_folder:
                    idx.doc_to_folder[doc_id] = title
                if doc_id not in idx.documents:
                    idx.documents[doc_id] = doc
        except Exception as e:
            log.warning(f"  could not fetch docs for folder '{title}': {e}")

    log.info(f"  indexed {len(idx.documents)} documents across {len(idx.id_to_title)} folders")

    # 3. Recently touched documents not in any folder ("Unfiled" in Granola)
    recent_docs = client.get_recent_documents()
    unfiled_added = 0
    for doc in recent_docs:
        if not isinstance(doc, dict):
            continue
        doc_id = doc.get("id")
        if not doc_id or doc_id in idx.documents:
            continue
        idx.documents[doc_id] = doc
        unfiled_added += 1
    if unfiled_added:
        log.info(f"  found {unfiled_added} additional unfiled document(s)")

    return idx

# ---------------------------------------------------------------------------
# Filename / folder name sanitization
# ---------------------------------------------------------------------------

_INVALID_RE = re.compile(r'[\\/:*?"<>|]+')
_WS_RE = re.compile(r"\s+")


def sanitize_segment(s: str, max_len: int = 120) -> str:
    s = _INVALID_RE.sub("-", s)
    s = _WS_RE.sub(" ", s).strip().strip(".")
    if len(s) > max_len:
        s = s[:max_len].rstrip()
    return s or "untitled"


def filename_for(doc: dict[str, Any]) -> str:
    raw_title = (doc.get("title") or "Untitled").strip() or "Untitled"
    created = doc.get("created_at") or doc.get("updated_at") or ""
    date = created[:10] if created else "0000-00-00"
    return f"{date}_{sanitize_segment(raw_title)}.docx"


def utc_now() -> str:
    return datetime.utcnow().isoformat(timespec="seconds") + "Z"


def atomic_write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(f".{path.name}.tmp")
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


def atomic_write_json(path: Path, data: Any) -> None:
    atomic_write_text(path, json.dumps(data, indent=2, sort_keys=True, ensure_ascii=False) + "\n")


def participant_names_for(doc: dict[str, Any]) -> list[str]:
    people = doc.get("people") or {}
    participant_names: list[str] = []
    if isinstance(people, dict):
        for grp in people.values():
            if isinstance(grp, list):
                for person in grp:
                    if isinstance(person, dict):
                        name = person.get("name") or person.get("email")
                        if name:
                            participant_names.append(str(name))
    return participant_names


# ---------------------------------------------------------------------------
# .docx builder
# ---------------------------------------------------------------------------

def render_html_to_docx(doc_obj, html_text: str) -> None:
    """
    Minimal HTML → docx renderer for Granola's panel format.
    Handles: h1-h6, p, ul/ol, li (with nesting), strong/b, em/i, br, a.
    """
    from html.parser import HTMLParser

    class _Parser(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.list_stack: list[str] = []
            self.current_para = None
            self.inline_stack: list[str] = []
            self.pending_heading: int | None = None

        def _ensure_para(self, style: str | None = None):
            if self.current_para is None:
                if style:
                    try:
                        self.current_para = doc_obj.add_paragraph(style=style)
                    except KeyError:
                        self.current_para = doc_obj.add_paragraph()
                else:
                    self.current_para = doc_obj.add_paragraph()
            return self.current_para

        def _flush_para(self) -> None:
            self.current_para = None

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._flush_para()
                level = min(int(tag[1]) + 1, 9)
                self.pending_heading = level
            elif tag == "p":
                self._flush_para()
            elif tag in ("ul", "ol"):
                self.list_stack.append(tag)
            elif tag == "li":
                self._flush_para()
                depth = len(self.list_stack)
                is_ordered = self.list_stack and self.list_stack[-1] == "ol"
                if is_ordered:
                    style = "List Number"
                else:
                    style = "List Bullet" if depth <= 1 else f"List Bullet {min(depth, 3)}"
                try:
                    self.current_para = doc_obj.add_paragraph(style=style)
                except KeyError:
                    self.current_para = doc_obj.add_paragraph(style="List Bullet")
            elif tag in ("strong", "b"):
                self.inline_stack.append("bold")
            elif tag in ("em", "i"):
                self.inline_stack.append("italic")
            elif tag == "br":
                if self.current_para:
                    self.current_para.add_run().add_break()

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._flush_para()
                self.pending_heading = None
            elif tag == "p":
                self._flush_para()
            elif tag in ("ul", "ol"):
                if self.list_stack:
                    self.list_stack.pop()
            elif tag == "li":
                self._flush_para()
            elif tag in ("strong", "b") and "bold" in self.inline_stack:
                self.inline_stack.remove("bold")
            elif tag in ("em", "i") and "italic" in self.inline_stack:
                self.inline_stack.remove("italic")

        def handle_data(self, data):
            if not data.strip() and self.current_para is None:
                return
            if self.pending_heading is not None:
                doc_obj.add_heading(data.strip(), level=self.pending_heading)
                self.pending_heading = None
                return
            p = self._ensure_para()
            run = p.add_run(data)
            if "bold" in self.inline_stack:
                run.bold = True
            if "italic" in self.inline_stack:
                run.italic = True

    _Parser().feed(html_text)


def render_markdown_to_docx(doc_obj, markdown_text: str) -> None:
    """Minimal markdown → docx renderer. Handles headings, bullets, bold."""
    if not markdown_text:
        doc_obj.add_paragraph("(no notes)")
        return

    for raw in markdown_text.replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            doc_obj.add_paragraph("")
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            doc_obj.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4) + 1)
            continue
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            style = "List Bullet" if indent == 0 else f"List Bullet {min(indent + 1, 3)}"
            try:
                p = doc_obj.add_paragraph(style=style)
            except KeyError:
                p = doc_obj.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, m.group(2))
            continue
        m = re.match(r"^(\s*)\d+[.)]\s+(.*)$", line)
        if m:
            try:
                p = doc_obj.add_paragraph(style="List Number")
            except KeyError:
                p = doc_obj.add_paragraph()
            _add_runs_with_bold(p, m.group(2))
            continue
        p = doc_obj.add_paragraph()
        _add_runs_with_bold(p, line)


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_bold(p, text: str) -> None:
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def build_docx(
    doc: dict[str, Any],
    folder_title: str,
    panels: list[dict[str, Any]],
    transcript: list[dict[str, Any]],
    output_path: Path,
) -> None:
    from docx import Document

    docx = Document()

    # Title
    title = doc.get("title") or "Untitled"
    docx.add_heading(title, level=0)

    # Subtitle: date · folder · participants
    created = doc.get("created_at") or ""
    date_str = ""
    if created:
        try:
            date_str = datetime.fromisoformat(created.replace("Z", "+00:00")).strftime("%b %d, %Y %I:%M %p")
        except ValueError:
            date_str = created[:19]

    participant_names = participant_names_for(doc)

    sub_parts = [p for p in (date_str, folder_title, ", ".join(participant_names[:6])) if p]
    if sub_parts:
        sub_para = docx.add_paragraph(" · ".join(sub_parts))
        sub_para.runs[0].italic = True

    docx.add_paragraph("─" * 40)

    # Notes section — panels (AI summaries) first, fallback to raw markdown
    docx.add_heading("Notes", level=1)
    rendered_any = False
    for panel in panels or []:
        content = (panel.get("original_content") or "").strip()
        if not content:
            continue
        panel_title = (panel.get("title") or "Summary").strip()
        docx.add_heading(panel_title, level=2)
        render_html_to_docx(docx, content)
        rendered_any = True

    if not rendered_any:
        notes_md = (doc.get("notes_markdown") or doc.get("notes_plain") or "").strip()
        if notes_md:
            render_markdown_to_docx(docx, notes_md)
        else:
            docx.add_paragraph("(no notes)")

    docx.add_paragraph("")
    docx.add_paragraph("─" * 40)

    # Transcript section
    docx.add_heading("Transcript", level=1)
    if not transcript:
        docx.add_paragraph("(no transcript available)")
    else:
        for seg in transcript:
            text = seg.get("text")
            if text is None or text == "":
                continue
            docx.add_paragraph(text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f".{output_path.name}.tmp")
    docx.save(tmp_path)
    tmp_path.replace(output_path)

# ---------------------------------------------------------------------------
# Manifest
# ---------------------------------------------------------------------------

def load_manifest() -> dict[str, dict[str, Any]]:
    if MANIFEST_PATH.exists():
        with MANIFEST_PATH.open() as f:
            return json.load(f)
    return {}


def save_manifest(manifest: dict[str, dict[str, Any]]) -> None:
    atomic_write_json(MANIFEST_PATH, manifest)

# ---------------------------------------------------------------------------
# Main sync logic
# ---------------------------------------------------------------------------

@dataclass
class SyncStats:
    created: int = 0
    updated: int = 0
    moved:   int = 0
    duplicates: int = 0
    skipped: int = 0
    failed:  int = 0


def manifest_entry_for(
    doc_id: str,
    doc: dict[str, Any],
    folder_safe: str,
    output_path: Path,
    notes_status: str,
    transcript_status: str,
    transcript_count: int,
    errors: list[str],
) -> dict[str, Any]:
    return {
        "document_id": doc_id,
        "updated_at": doc.get("updated_at") or "",
        "created_at": doc.get("created_at"),
        "folder": folder_safe,
        "filename": output_path.name,
        "title": doc.get("title"),
        "docx_path": str(output_path),
        "notes_status": notes_status,
        "transcript_status": transcript_status,
        "transcript_segment_count": transcript_count,
        "latest_error": "; ".join(errors) if errors else None,
        "synced_at": utc_now(),
    }


def run_sync(args: argparse.Namespace, log: logging.Logger) -> int:
    client = GranolaClient(log)

    log.info("Fetching folders and documents from Granola API")
    folder_index = build_folder_index(client, log)

    documents = list(folder_index.documents.values())
    log.info(f"Processing {len(documents)} documents across {len(folder_index.id_to_title)} folders")

    manifest = load_manifest()
    stats = SyncStats()

    BUILD_DIR.mkdir(exist_ok=True)

    folder_filter = args.folder.lower() if args.folder else None

    docs_to_process: list[tuple[str, dict, str]] = []
    for doc in documents:
        doc_id = doc.get("id")
        if not doc_id:
            continue
        if doc.get("deleted_at") or doc.get("was_trashed"):
            continue
        if not (doc.get("title") or doc.get("notes_markdown") or doc.get("notes_plain")):
            continue
        curr_folder = folder_index.doc_to_folder.get(doc_id, UNFILED_FOLDER)
        if folder_filter and folder_filter not in curr_folder.lower():
            continue
        docs_to_process.append((doc_id, doc, curr_folder))

    # Deduplicate by output filename — two Granola docs can share the same title+date.
    # Keep the one with the most recent updated_at (it's more likely to have AI panels).
    path_to_best: dict[str, tuple[str, dict, str]] = {}
    duplicate_items: list[tuple[str, dict, str, str]] = []
    for item in docs_to_process:
        did, doc, folder = item
        key = f"{sanitize_segment(folder)}/{filename_for(doc)}"
        existing = path_to_best.get(key)
        if existing is None:
            path_to_best[key] = item
        elif (doc.get("updated_at") or "") > (existing[1].get("updated_at") or ""):
            duplicate_items.append((existing[0], existing[1], existing[2], did))
            path_to_best[key] = item
        else:
            duplicate_items.append((did, doc, folder, existing[0]))
    if len(path_to_best) < len(docs_to_process):
        log.info(f"  found {len(docs_to_process) - len(path_to_best)} duplicate filename(s); suppressing duplicate docx")
    docs_to_process = list(path_to_best.values())

    log.info(f"  {len(docs_to_process)} documents to consider (after filters)")

    for doc_id, doc, curr_folder in docs_to_process:
        try:
            entry = manifest.get(doc_id, {})
            new_filename = filename_for(doc)
            curr_folder_safe = sanitize_segment(curr_folder)

            # 1. Handle reclassification (folder change) — move the local file
            prev_folder   = entry.get("folder")
            prev_filename = entry.get("filename")
            if prev_folder and prev_filename and (prev_folder != curr_folder_safe or prev_filename != new_filename):
                local_src = BUILD_DIR / prev_folder / prev_filename
                local_dst = BUILD_DIR / curr_folder_safe / new_filename
                if local_src != local_dst and local_src.exists():
                    log.info(f"  moving: {prev_folder}/{prev_filename} → {curr_folder_safe}/{new_filename}")
                    local_dst.parent.mkdir(parents=True, exist_ok=True)
                    if not args.dry_run:
                        shutil.move(str(local_src), str(local_dst))
                    stats.moved += 1

            # 2. Content change check
            curr_updated = doc.get("updated_at") or ""
            output_path = BUILD_DIR / curr_folder_safe / new_filename
            if (
                not args.full
                and entry.get("updated_at") == curr_updated
                and entry.get("folder") == curr_folder_safe
                and entry.get("filename") == new_filename
                and output_path.exists()
            ):
                stats.skipped += 1
                continue

            # 3. Fetch panels + transcript and build .docx
            log.info(f"  {'[dry-run] ' if args.dry_run else ''}↓ {curr_folder_safe}/{new_filename}")

            panels: list[dict] = []
            errors: list[str] = []
            notes_status = "ok"
            try:
                panels = client.get_document_panels(doc_id)
            except Exception as e:
                notes_status = "failed"
                error = f"panels fetch failed: {e}"
                errors.append(error)
                log.warning(f"  {error} for {doc_id}")

            transcript: list[dict] = []
            if not args.skip_transcripts:
                transcript_status = "ok"
                try:
                    transcript = client.get_transcript(doc_id)
                except Exception as e:
                    transcript_status = "failed"
                    error = f"transcript fetch failed: {e}"
                    errors.append(error)
                    log.warning(f"  {error} for {doc_id}")
            else:
                transcript_status = "skipped"

            if not args.dry_run:
                build_docx(doc, curr_folder, panels, transcript, output_path)

            if entry:
                stats.updated += 1
            else:
                stats.created += 1

            manifest[doc_id] = manifest_entry_for(
                doc_id=doc_id,
                doc=doc,
                folder_safe=curr_folder_safe,
                output_path=output_path,
                notes_status=notes_status,
                transcript_status=transcript_status,
                transcript_count=len(transcript),
                errors=errors,
            )
        except Exception as e:
            stats.failed += 1
            log.error(f"  failed for {doc_id}: {e}")

    for doc_id, doc, curr_folder, kept_doc_id in duplicate_items:
        try:
            entry = manifest.get(doc_id, {})
            new_filename = filename_for(doc)
            curr_folder_safe = sanitize_segment(curr_folder)
            curr_updated = doc.get("updated_at") or ""
            output_path = BUILD_DIR / curr_folder_safe / new_filename
            if (
                not args.full
                and entry.get("updated_at") == curr_updated
                and entry.get("folder") == curr_folder_safe
                and entry.get("filename") == new_filename
            ):
                stats.skipped += 1
                continue

            log.info(
                f"  {'[dry-run] ' if args.dry_run else ''}↳ suppress duplicate "
                f"{curr_folder_safe}/{new_filename} (docx kept from {kept_doc_id})"
            )
            errors = [f"duplicate filename; docx export suppressed in favor of {kept_doc_id}"]

            if entry:
                stats.updated += 1
            else:
                stats.created += 1
            stats.duplicates += 1

            manifest[doc_id] = manifest_entry_for(
                doc_id=doc_id,
                doc=doc,
                folder_safe=curr_folder_safe,
                output_path=output_path,
                notes_status="duplicate_suppressed",
                transcript_status="duplicate_suppressed",
                transcript_count=0,
                errors=errors,
            )
            manifest[doc_id]["duplicate_of"] = kept_doc_id
            manifest[doc_id]["docx_status"] = "duplicate_filename_suppressed"
        except Exception as e:
            stats.failed += 1
            log.error(f"  failed for duplicate {doc_id}: {e}")

    if not args.dry_run:
        save_manifest(manifest)

    log.info(
        f"DONE created={stats.created} updated={stats.updated} moved={stats.moved} "
        f"duplicates={stats.duplicates} "
        f"skipped={stats.skipped} failed={stats.failed}"
    )

    if not args.dry_run:
        _notify(stats)

    return 0 if stats.failed == 0 else 1


def _notify(stats: SyncStats) -> None:
    """Send a macOS notification summarising the sync result."""
    changed = stats.created + stats.updated + stats.moved
    if stats.failed:
        title   = "Granola Sync — errors"
        message = f"{stats.failed} doc(s) failed. Check logs for details."
    elif changed == 0:
        return  # nothing new — no notification needed
    else:
        parts = []
        if stats.created: parts.append(f"{stats.created} new")
        if stats.updated: parts.append(f"{stats.updated} updated")
        if stats.moved:   parts.append(f"{stats.moved} moved")
        title   = "Granola Sync — done"
        message = ", ".join(parts)

    try:
        subprocess.run(
            ["osascript", "-e",
             f'display notification "{message}" with title "{title}" sound name "Frog"'],
            check=False, capture_output=True,
        )
    except Exception:
        pass  # notifications are best-effort


def run_healthcheck(log: logging.Logger) -> int:
    checks: list[tuple[str, bool, str]] = []

    def record(name: str, ok: bool, detail: str = "") -> None:
        checks.append((name, ok, detail))
        status = "OK" if ok else "FAIL"
        log.info(f"healthcheck {status}: {name}{f' — {detail}' if detail else ''}")

    try:
        import docx  # noqa: F401
        record("python-docx import", True)
    except Exception as e:
        record("python-docx import", False, str(e))

    try:
        import cryptography  # noqa: F401
        record("cryptography import", True)
    except Exception as e:
        record("cryptography import", False, str(e))

    client = GranolaClient(log)
    try:
        token = client._token()
        record("Granola token", bool(token), "token loaded")
    except Exception as e:
        record("Granola token", False, str(e))

    try:
        metadata = client.get_all_document_lists_metadata()
        record("Granola API folders", isinstance(metadata, dict), f"{len(metadata)} folder metadata entries")
    except Exception as e:
        record("Granola API folders", False, str(e))

    for label, directory in (("docx output", BUILD_DIR), ("sync state", STATE_DIR)):
        try:
            directory.mkdir(parents=True, exist_ok=True)
            probe = directory / ".healthcheck.tmp"
            probe.write_text("ok\n", encoding="utf-8")
            probe.unlink()
            record(f"{label} writable", True, str(directory))
        except Exception as e:
            record(f"{label} writable", False, str(e))

    failed = [name for name, ok, _ in checks if not ok]
    if failed:
        log.error(f"healthcheck failed: {', '.join(failed)}")
        return 1
    log.info("healthcheck passed")
    return 0


def run_reconcile(args: argparse.Namespace, log: logging.Logger) -> int:
    client = GranolaClient(log)
    log.info("Reconciling Granola API, manifest, and docx output")
    folder_index = build_folder_index(client, log)
    manifest = load_manifest()

    api_docs = {
        doc_id: doc
        for doc_id, doc in folder_index.documents.items()
        if isinstance(doc, dict) and not doc.get("deleted_at") and not doc.get("was_trashed")
    }
    folder_filter = args.folder.lower() if args.folder else None
    if folder_filter:
        api_docs = {
            doc_id: doc
            for doc_id, doc in api_docs.items()
            if folder_filter in folder_index.doc_to_folder.get(doc_id, UNFILED_FOLDER).lower()
        }

    docx_files = list(BUILD_DIR.rglob("*.docx")) if BUILD_DIR.exists() else []
    manifest_docx_paths = set()
    missing_docx: list[str] = []
    missing_manifest: list[str] = []

    for doc_id, doc in api_docs.items():
        folder = sanitize_segment(folder_index.doc_to_folder.get(doc_id, UNFILED_FOLDER))
        docx_path = BUILD_DIR / folder / filename_for(doc)
        entry = manifest.get(doc_id, {})
        if isinstance(entry, dict) and entry.get("docx_status") == "duplicate_filename_suppressed":
            pass
        elif not docx_path.exists():
            missing_docx.append(doc_id)
        if doc_id not in manifest:
            missing_manifest.append(doc_id)

    for entry in manifest.values():
        if not isinstance(entry, dict):
            continue
        docx_path = entry.get("docx_path")
        if docx_path:
            manifest_docx_paths.add(Path(docx_path))
        elif entry.get("folder") and entry.get("filename"):
            manifest_docx_paths.add(BUILD_DIR / entry["folder"] / entry["filename"])

    orphan_manifest = sorted(set(manifest) - set(api_docs))
    untracked_docx = [p for p in docx_files if p not in manifest_docx_paths]

    log.info(f"  API documents: {len(api_docs)}")
    log.info(f"  manifest entries: {len(manifest)}")
    log.info(f"  docx files: {len(docx_files)}")
    log.info(f"  missing from manifest: {len(missing_manifest)}")
    log.info(f"  missing docx files: {len(missing_docx)}")
    log.info(f"  manifest entries not in API: {len(orphan_manifest)}")
    log.info(f"  docx files not tracked by manifest: {len(untracked_docx)}")

    for label, items in (
        ("missing_manifest", missing_manifest[:10]),
        ("missing_docx", missing_docx[:10]),
        ("orphan_manifest", orphan_manifest[:10]),
    ):
        if items:
            log.info(f"  sample {label}: {', '.join(items)}")
    if untracked_docx[:10]:
        log.info("  sample untracked_docx: " + "; ".join(str(p) for p in untracked_docx[:10]))
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Sync Granola notes to local .docx files")
    parser.add_argument("--dry-run",          action="store_true", help="Don't write files")
    parser.add_argument("--full",             action="store_true", help="Ignore manifest, rebuild everything")
    parser.add_argument("--folder",           type=str, default=None, help="Restrict to folder (substring match)")
    parser.add_argument("--skip-transcripts", action="store_true", help="Skip transcript fetches (faster)")
    parser.add_argument("--healthcheck",      action="store_true", help="Verify dependencies, auth, API, and write access")
    parser.add_argument("--reconcile",        action="store_true", help="Compare Granola API, manifest, and docx output")
    args = parser.parse_args()

    log = setup_logging()
    log.info("=" * 60)
    log.info(
        f"sync_granola starting — output: {BUILD_DIR} state: {STATE_DIR} "
        f"(dry_run={args.dry_run}, full={args.full}, folder={args.folder})"
    )
    try:
        if args.healthcheck:
            return run_healthcheck(log)
        if args.reconcile:
            return run_reconcile(args, log)
        return run_sync(args, log)
    except SystemExit:
        raise
    except Exception as e:
        log.exception(f"Unhandled error: {e}")
        return 2


if __name__ == "__main__":
    sys.exit(main())
